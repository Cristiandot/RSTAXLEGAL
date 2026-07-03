# -*- coding: utf-8 -*-
"""
Genera las plantillas docxtemplater del CONTRATO DE TRABAJADOR DE CASA PARTICULAR
PUERTAS AFUERA (plazo fijo e indefinido). Placeholders {TAG} rellenados por el
apartado Casa Particular del panel (app/(app)/casa-particular/actions.ts).
Formato casa: titulo + subtitulo centrados, cuerpo justificado, Arial 11.

Correr desde app/:  py scripts/gen_plantilla_casa_particular.py
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT = "Arial"


def build(indefinido: bool):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)

    def titulo(text, size=14):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(size)
        return p

    def cuerpo(runs, justify=True, space_before=6):
        """runs: lista de (texto, bold). Cada item es UN run (placeholders intactos)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(space_before)
        for texto, bold in runs:
            r = p.add_run(texto)
            r.bold = bold
            r.font.name = FONT
            r.font.size = Pt(11)
        return p

    def clausula(label, texto):
        if texto:
            cuerpo([(label + " ", True), (texto, False)], space_before=10)
        else:
            cuerpo([(label, True)], space_before=10)

    # ---- Titulo ----
    titulo("CONTRATO DE TRABAJO", 14)
    titulo("Trabajador de Casa Particular Puertas Afuera", 12)

    # ---- Preambulo ----
    cuerpo([
        ("En ", False), ("{CIUDAD_FIRMA}", False), (", a ", False),
        ("{FECHA_INICIO_CONTRATO}", False),
        (", entre ", False), ("{RAZON_SOCIAL}", False), (", RUT ", False),
        ("{RUT_EMPRESA}", False), (", con domicilio en ", False),
        ("{DIRECCION_EMPRESA}", False),
        (", en adelante “el Empleador”, por una parte, y doña ", False),
        ("{NOMBRE_EMPLEADO}", False), (", Cédula de Identidad N° ", False),
        ("{RUT_EMPLEADO}", False), (", fecha de nacimiento ", False),
        ("{FECHA_NACIMIENTO}", False), (", con domicilio en ", False),
        ("{DIRECCION_EMPLEADO}", False), (", comuna de ", False),
        ("{COMUNA_EMPLEADO}", False), (", de nacionalidad ", False),
        ("{NACIONALIDAD_EMPLEADO}", False),
        (", en adelante “la trabajadora”, se conviene un contrato de trabajo "
         "cuyas cláusulas son las siguientes:", False),
    ], space_before=12)

    # ---- PRIMERO ----
    clausula("PRIMERO: De la naturaleza de los servicios.",
             "La trabajadora se obliga a desempeñar trabajos de asistencia propios o "
             "inherentes al hogar, de acuerdo a las instrucciones que al efecto sean "
             "impartidas por el Empleador. La trabajadora queda obligada a cumplir leal y "
             "correctamente con todos los deberes que le imponga este instrumento o aquellos "
             "que se deriven de las labores contratadas, empleando para ello la mayor "
             "diligencia y dedicación.")
    cuerpo([
        ("Para una adecuada descripción de las obligaciones que derivan del presente "
         "contrato, las partes acuerdan que la trabajadora desempeñará de forma "
         "principal las siguientes tareas: ", False),
        ("{FUNCIONES_CASA}", False), (".", False),
    ])

    # ---- SEGUNDO ----
    clausula("SEGUNDO: Del lugar de prestación de los servicios.", "")
    cuerpo([
        ("Las partes acuerdan que los servicios deberán ser prestados en el domicilio "
         "ubicado en ", False), ("{DIRECCION_EMPRESA}", False),
        (", sin perjuicio de otras ubicaciones que, previo acuerdo y consentimiento de la "
         "trabajadora, podrá, de manera ocasional, desarrollar sus labores.", False),
    ])

    # ---- TERCERO ----
    clausula("TERCERO: Del monto, forma y período de pago de las remuneraciones.",
             "La trabajadora tendrá derecho a percibir las siguientes prestaciones a "
             "título de remuneración:")
    cuerpo([("• Sueldo ascendente a ", False), ("{SUELDO_BASE}", False), (".", False)],
           justify=False)
    cuerpo([("• Asignación de movilización de $", False),
            ("{ASIGNACION_MOVILIZACION}", False), (" por día trabajado.", False)],
           justify=False)
    cuerpo([
        ("Las remuneraciones se pagarán por períodos mensuales vencidos, el "
         "último día hábil de cada mes. De las sumas anteriores se "
         "deducirán los impuestos que las graven, las cotizaciones de seguridad social "
         "y las demás que correspondan.", False),
    ])
    cuerpo([
        ("El pago de la remuneración se hará en dinero efectivo. A solicitud de la "
         "trabajadora, el pago podrá realizarse por medio de cheque, o depósito en "
         "cuenta vista o cuenta corriente que la trabajadora indique.", False),
    ])

    # ---- CUARTO ----
    clausula("CUARTO: Cotizaciones de Seguridad Social.",
             "De la remuneración bruta que tiene derecho a percibir la trabajadora en el "
             "respectivo mes, el Empleador deberá practicar los descuentos legales para el "
             "pago de cotizaciones de seguridad social (cotizaciones para pensiones y "
             "cotizaciones para salud).")
    cuerpo([
        ("Asimismo, el Empleador se obliga a enterar mensualmente, en la A.F.P. que la "
         "trabajadora determine, un 4,11% de su remuneración imponible, por el tiempo de "
         "duración del contrato, plazo que no podrá exceder de 11 años a contar "
         "de la fecha de inicio de la relación laboral, con el objeto de financiar la "
         "indemnización a todo evento por término de contrato a que tiene derecho la "
         "trabajadora de casa particular.", False),
    ])
    cuerpo([("Al efecto, la trabajadora declara encontrarse afiliada a las siguientes "
             "instituciones de seguridad social:", False)])
    cuerpo([("• A.F.P.: ", False), ("{PREVISION}", False), (".", False)], justify=False)
    cuerpo([("• Salud: ", False), ("{INSTITUCION_SALUD}", False), (".", False)], justify=False)
    cuerpo([
        ("Será de cargo del Empleador declarar y pagar la cotización para el seguro "
         "de accidentes del trabajo y enfermedades profesionales de la Ley N° 16.744, "
         "así como las cotizaciones del seguro de cesantía de la Ley N° 19.728.", False),
    ])

    # ---- QUINTO ----
    clausula("QUINTO: Jornada de Trabajo.", "")
    cuerpo([
        ("La trabajadora estará sujeta a una jornada ordinaria de ", False),
        ("{HORAS_SEMANALES}", False), (" horas semanales, ", False),
        ("{DISTRIBUCION_JORNADA}", False), (".", False),
    ])

    # ---- SEXTO ----
    clausula("SEXTO: Fecha de inicio y duración del contrato.", "")
    if indefinido:
        cuerpo([
            ("Las partes dejan constancia que la trabajadora comenzó a prestar servicios "
             "para el Empleador con fecha ", False), ("{FECHA_INICIO_CONTRATO}", False),
            (". El presente contrato es de duración indefinida. ", False),
            ("{CLAUSULAS_ADICIONALES}", False),
        ])
    else:
        cuerpo([
            ("Las partes dejan constancia que la trabajadora comenzará a prestar servicios "
             "para el Empleador con fecha ", False), ("{FECHA_INICIO_CONTRATO}", False),
            (". El presente contrato tendrá vigencia hasta el ", False),
            ("{FECHA_TERMINO_CONTRATO}", False), (". ", False),
            ("{CLAUSULAS_ADICIONALES}", False),
        ])

    # ---- SEPTIMO ----
    clausula("SÉPTIMO:",
             "El presente contrato de trabajo se firma por la trabajadora y el Empleador en dos "
             "ejemplares del mismo tenor y fecha, quedando en este mismo acto un ejemplar en "
             "poder de cada contratante.")

    # ---- OCTAVO ----
    clausula("OCTAVO:",
             "Las partes manifiestan estar en conocimiento de que se encuentra prohibido exigir "
             "a la trabajadora el uso obligatorio de uniforme, delantal o cualquier otro "
             "distintivo o vestimenta identificadores en espacios, lugares o establecimientos "
             "públicos como parques, plazas, playas, restaurantes, hoteles, locales "
             "comerciales, clubes sociales y otros de similar naturaleza.")

    # ---- Firmas ----
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2)

    def cel(row, col, text, bold=False):
        c = tbl.cell(row, col)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.name = FONT
        r.font.size = Pt(11)

    cel(0, 0, "_____________________________")
    cel(0, 1, "_____________________________")
    cel(1, 0, "{RAZON_SOCIAL}", bold=True)
    cel(1, 1, "{NOMBRE_EMPLEADO}", bold=True)
    cel(2, 0, "RUT {RUT_EMPRESA}")
    cel(2, 1, "C.I. N° {RUT_EMPLEADO}")
    cel(3, 0, "El Empleador")
    cel(3, 1, "La trabajadora")

    out_dir = os.path.join("plantillas", "GENERICO")
    os.makedirs(out_dir, exist_ok=True)
    sufijo = "Indefinido" if indefinido else "Plazo Fijo"
    out = os.path.join(out_dir, f"CONTRATO Casa Particular Puertas Afuera - {sufijo}.docx")
    doc.save(out)
    print("OK:", out)


build(indefinido=False)
build(indefinido=True)
