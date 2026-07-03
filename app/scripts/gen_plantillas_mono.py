# -*- coding: utf-8 -*-
"""
Generador de plantillas de contrato de MONO NEGOCIOS SpA (77.096.809-7).
Base: contrato tipo "Part Time" de atención de cafetería entregado por
Cristian (jul-2026), local Av. Ignacio Carrera Pinto N° 184, Reñaca.
La jornada usa placeholders ({HORAS_SEMANALES} / {DISTRIBUCION_JORNADA}) y la
duración {TIPO_CONTRATO}: una plantilla sirve PT/FT y plazo fijo/indefinido.

Variantes:
- Atención del Local: jornada por {DISTRIBUCION_JORNADA} (texto libre).
- Garzón: jornada en turnos rotativos con las TABLAS de turnos referenciales
  (mismo catálogo de JEREZ — se lee de gen_plantillas_jerez.py, fuente única).

Ejecutar:  py scripts/gen_plantillas_mono.py   (desde app/)
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "plantillas/MONO NEGOCIOS 77096809-7/"
OUT_ATENCION = BASE + "CONTRATO Atencion del Local (PT-FT).docx"
OUT_ATENCION_EXT = BASE + "CONTRATO Atencion del Local Extranjero (PT-FT).docx"
OUT_GARZON = BASE + "CONTRATO Garzon (PT-FT).docx"
OUT_GARZON_EXT = BASE + "CONTRATO Garzon Extranjero (PT-FT).docx"

# Cláusula obligatoria en toda plantilla de nacionalidad extranjero (regla RSTL).
CLAUSULA_EXTRANJERO = [
    ("head", "Se establecen las siguientes:"),
    ("li", "- Vigencia: La obligación de prestar servicios emanada del presente contrato sólo podrá cumplirse una vez que el trabajador haya obtenido la visación de residencia correspondiente en Chile o el permiso especial de trabajo para extranjeros con visa en trámite."),
    ("li", "- Régimen previsional: El empleador se compromete a efectuar las retenciones correspondientes y entregarlas a las instituciones de seguridad social, salvo que las partes se acojan a la Ley 18.156."),
    ("li", "- Impuesto a la renta: El empleador se obliga a responder por el pago del impuesto a la renta correspondiente a la remuneración del trabajador extranjero, para rentas superiores a 13,5 UTM."),
]

# --- Turnos referenciales: mismo catálogo de JEREZ (fuente única en
# gen_plantillas_jerez.py; se extrae el literal sin importar el módulo porque
# ese script genera archivos al importarse). ---
_JEREZ = os.path.join(os.path.dirname(os.path.abspath(__file__)), "gen_plantillas_jerez.py")
with open(_JEREZ, encoding="utf-8") as f:
    _src = f.read()
TURNOS_RAW = _src.split('TURNOS_RAW = """', 1)[1].split('"""', 1)[0]

CIERRE_TURNOS = (
    "Las jornadas y turnos señalados en el presente contrato tendrán carácter referencial. "
    "Además, deben ser pactados expresamente y comunicados al trabajador con una anticipación mínima de una semana, "
    "y podrán ajustarse según las necesidades operativas de la empresa, respetando siempre los límites legales de "
    "jornada y descanso establecidos en la legislación vigente. En todo caso, dichas modificaciones no podrán implicar "
    "menoscabo para el trabajador ni exceder los límites legales aplicables. Se deja establecido que los horarios de "
    "los turnos se encuentran en el Reglamento Interno de la empresa."
)


def parse_turnos(raw):
    """[(categoria, [(subgrupo, [(codigo, horario)])])]"""
    cats = []
    for line in raw.strip().splitlines():
        t = line.strip()
        if not t:
            continue
        low = t.lower()
        if low.startswith("jornada de"):
            cats.append((t.rstrip("."), []))
        elif low.startswith("turnos am") or low.startswith("turnos pm"):
            cats[-1][1].append((t.rstrip("."), []))
        elif t.startswith("●"):
            body = t[1:].strip()
            if body.lower().startswith("turno"):
                resto = body[5:].strip()
                cod, _, hor = resto.partition(":")
                cats[-1][1][-1][1].append((cod.strip(), hor.strip()))
    return cats


def add_turnos(d):
    """Inserta los turnos como tablas (Turno | Horario) por categoría de horas."""
    def set_cell(cell, text, bold=False, size=9, shade=None):
        cell.paragraphs[0].alignment = AL.LEFT
        r = cell.paragraphs[0].add_run(text); r.bold = bold; r.font.size = Pt(size)
        if shade:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), shade)
            cell._tc.get_or_add_tcPr().append(sh)
    for cat_title, groups in parse_turnos(TURNOS_RAW):
        h = d.add_paragraph(); h.paragraph_format.space_before = Pt(6); h.paragraph_format.keep_with_next = True
        h.add_run(cat_title).bold = True
        tbl = d.add_table(rows=0, cols=2); tbl.style = "Table Grid"; tbl.autofit = False
        for sub_title, turnos in groups:
            c = tbl.add_row().cells; c[0].merge(c[1])
            set_cell(c[0], sub_title, bold=True, shade="EFEFEF")
            for cod, hor in turnos:
                c = tbl.add_row().cells
                set_cell(c[0], cod, bold=True); set_cell(c[1], hor)
                c[0].width = Cm(1.6); c[1].width = Cm(15.4)
        d.add_paragraph()


def build_contrato(out_path: str, con_turnos: bool, extranjero: bool = False):
    d = Document()
    d.styles["Normal"].paragraph_format.space_after = Pt(0)

    def cab(t):
        p = d.add_paragraph(); p.alignment = AL.CENTER; r = p.add_run(t); r.bold = True

    def vacio():
        d.add_paragraph()

    def cl(lead, resto=""):
        p = d.add_paragraph(); p.alignment = AL.JUSTIFY
        if lead:
            r = p.add_run(lead); r.bold = True
        if resto:
            p.add_run(resto)

    def num(t, bold=False):
        # ítem numerado (el número va en el texto), con sangría colgante
        p = d.add_paragraph(); p.alignment = AL.JUSTIFY
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(t); r.bold = bold

    def bullet(t):
        # punteo real de Word (estilo List Bullet), segundo nivel
        p = d.add_paragraph(style="List Bullet"); p.alignment = AL.JUSTIFY
        p.paragraph_format.left_indent = Pt(36)
        p.add_run(t)

    cab("CONTRATO INDIVIDUAL DE TRABAJO")
    cab("{CARGO}")
    vacio()
    cl("", 'En {CIUDAD_FIRMA}, a {FECHA_INICIO_CONTRATO}, entre {RAZON_SOCIAL}, RUT {RUT_EMPRESA}, representada por don {NOMBRE_REP_LEGAL}, Rut: {RUT_REP_LEGAL}, ambos domiciliados en {DIRECCION_EMPRESA}, comuna de {COMUNA_EMPRESA}, correo electrónico {EMAIL_EMPRESA}, que en adelante se denominará "el empleador", y don(ña) {NOMBRE_EMPLEADO}, de nacionalidad {NACIONALIDAD_EMPLEADO}, Rut {RUT_EMPLEADO}, fecha de nacimiento {FECHA_NACIMIENTO}, de estado civil {ESTADO_CIVIL}, con domicilio {DIRECCION_EMPLEADO}, comuna de {COMUNA_EMPLEADO}, correo electrónico {EMAILPERSONAL_EMPLEADO}, que en adelante se denominará "el trabajador", se ha convenido en el siguiente contrato de trabajo:')
    vacio()
    cl("PRIMERO: Naturaleza de los Servicios. ", "El Trabajador se compromete y obliga a realizar el trabajo como {CARGO}, desempeñando labores en el área indicada por el empleador y otras del giro, las cuales han de prestarse en las dependencias del empleador, ubicadas en {DIRECCION_EMPRESA}, comuna de {COMUNA_EMPRESA}. El trabajador podrá ser trasladado a otro domicilio o labores similares, dentro de la ciudad, por causa justificada, sin que ello importe menoscabo para el trabajador.")
    cl("Funciones, Responsabilidades, Obligaciones y Prohibiciones Específicas:")
    cl("A. Funciones")
    num("1. Atención al cliente y servicio a la mesa.", bold=True)
    bullet("Recepción cordial de clientes, entrega de cartas/menús y asesoría sobre los productos de la cafetería.")
    bullet("Toma de pedidos de manera ágil y exacta, ingresándolas en el sistema de comanda correspondiente.")
    bullet("Servicio de alimentos y bebidas (café, pastelería, sándwiches, etc.) en mesa o barra, cumpliendo con los estándares de presentación de la Empresa.")
    num("2. Operación y Caja.", bold=True)
    bullet("Entrega de la cuenta y cobro de los servicios a través de los medios de pago disponibles (efectivo, tarjetas, transferencias).")
    bullet("Apoyo en la barra para la preparación básica de pedidos rápidos o despacho de productos envasados si la operación lo requiere.")
    num("3. Higiene y Orden.", bold=True)
    bullet("Montaje, desmontaje y limpieza profunda de mesas y sillas inmediatamente después de su uso.")
    bullet("Mantención del aseo general del salón, barra y estaciones de trabajo, asegurando el cumplimiento de las normas sanitarias vigentes.")
    bullet("Orden y reposición de insumos (servilletas, azúcar, cubiertos, vitrinas) antes, durante y al cierre del turno.")
    num("4. Funciones conexas: ", bold=True)
    num("El Trabajador acepta colaborar en tareas conexas ante necesidades operativas, tales como apoyo en inventarios generales, cierre de local y verificación de aseo.")
    cl("B. Obligaciones.")
    num("1. Llegar puntualmente a su lugar de trabajo.")
    num("2. Realizar personalmente la labor convenida.")
    num("3. Efectuar el trabajo de acuerdo con las órdenes e instrucciones que emanen de la administración de la empresa, como, asimismo, las recibidas de sus jefes directos.")
    num("4. Ser respetuoso con sus superiores y compañeros de trabajo, y observar las órdenes que los primeros impartan en orden al buen servicio y/o los intereses de la empresa.")
    num("5. Registrar diariamente su hora de entrada y salida en el Libro de Asistencia o mecanismo afín. Se considerará falta grave que el trabajador registre indebidamente su asistencia en el Libro o que éste sea firmado por otro u otros funcionarios.")
    num("6. Observar, en todo momento, una conducta correcta y honorable, y desempeñar sus funciones con dignidad y responsabilidad.")
    num("7. Cumplir íntegramente la jornada diaria y semanal de trabajo a la que se encuentre afecto.")
    num("8. Informar, dentro de los sesenta minutos siguientes al inicio de sus labores, cualquier impedimento que lo haya imposibilitado de concurrir, justificando su inasistencia dentro de las cuarenta y ocho horas a partir de ese momento, con la respectiva licencia médica.")
    num("9. Dar aviso de inmediato de cualquier desperfecto que se detecte en los elementos de trabajo o instalaciones de la empresa, a efecto de evitar daños mayores.")
    num("10. Mantener en orden, higiene y aseo el lugar donde desempeñe sus labores.")
    num("11. Mantener siempre una excelente presencia y aspecto personal.")
    num("12. Cuidar con esmero, máquinas, implementos, equipos, instalaciones y útiles que se le entreguen para el desempeño de sus labores, los cuales deberá dejar limpios y/o guardados al término de su jornada. Ante cualquier destrucción, estos serán reparados y su arreglo o reemplazo será de cargo del trabajador. Esto no procederá cuando la destrucción de ellos sea a causa de su normal uso y desgaste.")
    num("13. Proceder a devolver todos los elementos puestos a su servicio, antes de cursar la firma del respectivo finiquito legal.")
    cl("C. Prohibiciones.")
    num("1. Consumo de Alcohol: Ingerir alcohol del inventario del local durante el turno o presentarse bajo sus efectos.")
    num("2. Descuadre de Caja: Diferencias injustificadas en la recaudación bajo su custodia.")
    num("3. Trato al Cliente: Discusiones, faltas de respeto o negligencia que afecte la imagen del local.")
    num("4. Seguridad: Dejar el local abierto o sin alarma de cierre, o, de manera negligente, se permita sustracciones de insumos, dinero o productos del local.")
    cl("", "El incumplimiento de alguna de las disposiciones indicadas, así como el incurrir en alguna de las prohibiciones, constituye incumplimiento grave de las obligaciones que le impone el contrato, siendo causal suficiente para poner término inmediato a éste.")
    vacio()
    if con_turnos:
        cl("SEGUNDO: ", "El trabajador cumplirá una jornada de trabajo de {HORAS_SEMANALES} horas a la semana, distribuida en turnos rotativos y variables que el empleador publicará semanalmente, con los descansos legales correspondientes. Con todo, el trabajador destinará una hora de colación diaria, no imputable a la jornada de trabajo. Para todos los efectos, los turnos referenciales de la empresa son los siguientes:")
        vacio()
        add_turnos(d)
        cl("", CIERRE_TURNOS)
    else:
        cl("SEGUNDO: ", "El trabajador cumplirá una jornada de trabajo de {HORAS_SEMANALES} horas a la semana, distribuida de la siguiente forma: {DISTRIBUCION_JORNADA}. Sin perjuicio de lo anterior, el empleador podrá modificar la distribución y extensión de los turnos en función de las necesidades operacionales de la Empresa, siempre respetando el máximo semanal permitido. Con todo, el trabajador destinará una hora de colación no imputable a la jornada de trabajo.")
    vacio()
    cl("TERCERO: ", "El Empleador se compromete a remunerar al Trabajador en la forma que se indica:")
    num("1) Sueldo Base: $ {SUELDO_BASE} ({SUELDO_BASE_PALABRA}).")
    num("2) Asignación de Pérdida de Caja: {ASIGNACION_CAJA_TEXTO}.")
    num("3) Asignación de movilización: $ {ASIGNACION_MOVILIZACION} ({ASIGNACION_MOVILIZACION_PALABRA}).")
    num("4) Asignación de Colación: $ {ASIGNACION_COLACION} ({ASIGNACION_COLACION_PALABRA}).")
    vacio()
    cl("CUARTO: ", "{GRATIFICACION_TEXTO}. Las remuneraciones antes referidas quedarán afectas a los descuentos previsionales y tributarios que correspondan y a los demás que autorice la ley. El trabajador acepta y autoriza, desde ya, que pueda descontarse de su remuneración el tiempo no trabajado, sea por permisos, atrasos, faltas, rotura de implementos, u otros motivos, como también por anticipos de remuneración solicitados por él.")
    vacio()
    cl("QUINTO: Confidencialidad. ", "Prohibición absoluta de divulgar recetas de coctelería, bases de datos de clientes o cifras de venta del local. El trabajador se obliga a no entregar, revelar, divulgar ni comunicar a ninguna persona, salvo a sus superiores jerárquicos o a la administración de {RAZON_SOCIAL}, cualquier información, antecedente, cifra o dato relativos a las actividades comerciales y/o financieras de {RAZON_SOCIAL} y/o de empresas relacionadas comercialmente a ésta (clientes, proveedores, relacionados comerciales, etc.), a las cuales pudiere tener acceso de acuerdo al desarrollo de su cargo. La infracción a esta norma supone incumplimiento grave de las obligaciones que impone el contrato.")
    vacio()
    cl("SEXTO: ", "La duración de este contrato será {TIPO_CONTRATO} ({FECHA_TERMINO_CONTRATO}), rigiendo para su término las disposiciones establecidas en el Código del Trabajo.")
    vacio()
    cl("SÉPTIMO: ", "Se entienden incorporadas al presente contrato todas las disposiciones legales que se dicten con posterioridad a la fecha de suscripción y que tengan relación con él.")
    vacio()
    cl("OCTAVO: ", "Se deja constancia que el trabajador ingresó a prestar servicios para el empleador con fecha {FECHA_INICIO_CONTRATO}.")
    vacio()
    cl("NOVENO: ", "El trabajador podrá hacer uso de su feriado legal anual cuando le corresponda según la ley. Sin embargo, el empleador se reserva el derecho de establecer la época en que ello resulte conveniente, habida consideración de las necesidades del servicio de su giro.")
    vacio()
    cl("DÉCIMO: ", "Toda modificación que se introduzca al presente contrato o a su anexo deberá constar por escrito, a su dorso o en una hoja de prolongación colocada a continuación, y deberá estar firmada por el trabajador y el empleador.")
    vacio()
    cl("DÉCIMO PRIMERO: ", "Las cartas de amonestación son el medio utilizado por el empleador para comunicar al trabajador los incumplimientos en que ha incurrido, los cuales, de persistir, dan derecho al empleador para poner término al contrato de trabajo por incumplimiento grave de las obligaciones que impone el contrato o la causal que corresponda.")
    vacio()
    if extranjero:
        cl("DÉCIMO SEGUNDO: CLÁUSULAS ESPECIALES (TRABAJADOR EXTRANJERO). ", CLAUSULA_EXTRANJERO[0][1])
        for _, texto in CLAUSULA_EXTRANJERO[1:]:
            num(texto)
        vacio()
        n_jur, n_ej = "DÉCIMO TERCERO", "DÉCIMO CUARTO"
    else:
        n_jur, n_ej = "DÉCIMO SEGUNDO", "DÉCIMO TERCERO"
    cl(f"{n_jur}: ", "Para todos los efectos legales derivados de este contrato, las partes fijan su domicilio en la ciudad y comuna de Viña del Mar, sometiéndose a la jurisdicción y competencia de sus Tribunales de Justicia.")
    vacio()
    cl(f"{n_ej}: ", "Se suscribe este instrumento en dos ejemplares de igual tenor, quedando uno de ellos en poder del empleador y el restante en poder del trabajador, quien declara recibirlo en este acto.")
    vacio(); vacio()
    tb = d.add_table(rows=3, cols=2); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    filas = [("{RAZON_SOCIAL}", "{NOMBRE_EMPLEADO}"),
             ("RUT {RUT_EMPRESA}", "RUT {RUT_EMPLEADO}"),
             ("EMPLEADOR", "TRABAJADOR")]
    for i, (a, b) in enumerate(filas):
        c = tb.rows[i].cells
        for cell, txt in ((c[0], a), (c[1], b)):
            cell.paragraphs[0].alignment = AL.CENTER
            cell.paragraphs[0].add_run(txt)
    d.save(out_path)
    print("Generado:", out_path, "| párrafos:", len(d.paragraphs))


if __name__ == "__main__":
    os.makedirs(BASE, exist_ok=True)
    build_contrato(OUT_ATENCION, con_turnos=False)
    build_contrato(OUT_ATENCION_EXT, con_turnos=False, extranjero=True)
    build_contrato(OUT_GARZON, con_turnos=True)
    build_contrato(OUT_GARZON_EXT, con_turnos=True, extranjero=True)
