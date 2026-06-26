# -*- coding: utf-8 -*-
"""
Generador de plantillas de contrato de JEREZ DE LA FRONTERA SPA.
Fuente única: incluye los TURNOS de trabajo (verbatim de Cristian) en la
cláusula SEGUNDO de TODOS los contratos de JEREZ, con su texto de cierre
referencial. Reusar para nuevos cargos/variantes.
Ejecutar:  py scripts/gen_plantillas_jerez.py   (desde app/)
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_turnos(raw):
    """[(categoria, [(subgrupo, [(codigo, horario)])])]"""
    cats = []
    for line in raw.strip().splitlines():
        t = line.strip()
        if not t:
            continue
        low = t.lower()
        if low.startswith("jornada de"):
            cats.append((t.rstrip("."), []))
        elif low.startswith("turnos am") or low.startswith("turnos pm"):
            cats[-1][1].append((t.rstrip("."), []))
        elif t.startswith("●"):
            body = t[1:].strip()
            if body.lower().startswith("turno"):
                resto = body[5:].strip()
                cod, _, hor = resto.partition(":")
                cats[-1][1][-1][1].append((cod.strip(), hor.strip()))
    return cats


def add_turnos(d):
    """Inserta los turnos como tablas (Turno | Horario) por categoría de horas."""
    def set_cell(cell, text, bold=False, size=9, shade=None):
        cell.paragraphs[0].alignment = AL.LEFT
        r = cell.paragraphs[0].add_run(text); r.bold = bold; r.font.size = Pt(size)
        if shade:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), shade)
            cell._tc.get_or_add_tcPr().append(sh)
    for cat_title, groups in parse_turnos(TURNOS_RAW):
        h = d.add_paragraph(); h.paragraph_format.space_before = Pt(6); h.paragraph_format.keep_with_next = True
        h.add_run(cat_title).bold = True
        tbl = d.add_table(rows=0, cols=2); tbl.style = "Table Grid"; tbl.autofit = False
        for sub_title, turnos in groups:
            c = tbl.add_row().cells; c[0].merge(c[1])
            set_cell(c[0], sub_title, bold=True, shade="EFEFEF")
            for cod, hor in turnos:
                c = tbl.add_row().cells
                set_cell(c[0], cod, bold=True); set_cell(c[1], hor)
                c[0].width = Cm(1.6); c[1].width = Cm(15.4)
        d.add_paragraph()

# --- TURNOS (verbatim) ---
TURNOS_RAW = """
JORNADA DE 7 HORAS APROXIMADAS.
TURNOS AM.
● Turno A1: Lunes a viernes 09:00 – 16:00, sábado 09:00 – 16:00. Domingo libre.
● Turno A2: Lunes a viernes 09:30 – 16:30, sábado 09:30 – 16:30. Domingo libre.
● Turno A3: Lunes a viernes 10:00 – 17:00, sábado 10:00 – 17:00. Domingo libre.
● Turno A4: Lunes a viernes 10:30 – 17:30, sábado 10:30 – 17:30. Domingo libre.
● Turno A5: Lunes a viernes 11:00 – 18:00, sábado 11:00 – 18:00. Domingo libre.
● Turno A6: Lunes a viernes 11:30 – 18:30, sábado 11:30 – 18:30. Domingo libre.
● Turno A7: Lunes a viernes 12:00 – 19:00, sábado 12:00 – 19:00. Domingo libre.
TURNOS PM.
● Turno P1: Lunes a viernes 15:00 – 22:00, sábado 15:00 – 22:00. Domingo libre.
● Turno P2: Lunes a viernes 15:30 – 22:30, sábado 15:30 – 22:30. Domingo libre.
● Turno P3: Lunes a viernes 16:00 – 23:00, sábado 16:00 – 23:00. Domingo libre.
● Turno P4: Lunes a viernes 16:30 – 23:30, sábado 16:30 – 23:30. Domingo libre.
● Turno P5: Lunes a viernes 13:00 – 20:00, sábado 13:00 – 20:00. Domingo libre.
● Turno P6: Lunes a viernes 13:30 – 20:30, sábado 13:30 – 20:30. Domingo libre.
● Turno P7: Lunes a viernes 14:00 – 21:00, sábado 14:00 – 21:00. Domingo libre.
● Turno N1: Lunes a viernes 17:00 – 00:00, sábado 17:00 – 00:00. Domingo libre.
● Turno N2: Lunes a viernes 17:30 – 00:30, sábado 17:30 – 00:30. Domingo libre.
● Turno N3: Lunes a viernes 18:00 – 01:00, sábado 18:00 – 01:00. Domingo libre.
JORNADA DE 7,5 HORAS APROXIMADAS
TURNOS AM
● Turno A1: Lunes a viernes 09:00 – 16:15, sábado 09:00 – 15:45. Domingo libre.
● Turno A2: Lunes a viernes 09:30 – 16:45, sábado 09:30 – 16:15. Domingo libre.
● Turno A3: Lunes a viernes 10:00 – 17:15, sábado 10:00 – 16:45. Domingo libre.
● Turno A4: Lunes a viernes 10:30 – 17:45, sábado 10:30 – 17:15. Domingo libre.
● Turno A5: Lunes a viernes 11:00 – 18:15, sábado 11:00 – 17:45. Domingo libre.
● Turno A6: Lunes a viernes 11:30 – 18:45, sábado 11:30 – 18:15. Domingo libre.
● Turno A7: Lunes a viernes 12:00 – 19:15, sábado 12:00 – 18:45. Domingo libre.
TURNOS PM
● Turno P1: Lunes a viernes 15:00 – 22:15, sábado 15:00 – 21:45. Domingo libre.
● Turno P2: Lunes a viernes 15:30 – 22:45, sábado 15:30 – 22:15. Domingo libre.
● Turno P3: Lunes a viernes 16:00 – 23:15, sábado 16:00 – 22:45. Domingo libre.
● Turno P4: Lunes a viernes 16:30 – 23:45, sábado 16:30 – 23:15. Domingo libre.
● Turno N1: Lunes a viernes 17:00 – 00:15, sábado 17:00 – 23:45. Domingo libre.
● Turno N2: Lunes a viernes 17:15 – 00:30, sábado 17:15 – 23:30. Domingo libre.
● Turno N3: Lunes a viernes 17:00 – 00:30, sábado 18:00 – 23:30. Domingo libre.
● Turno N4: Lunes a viernes 17:15 – 00:30, sábado 17:45 – 23:45. Domingo libre.
JORNADA DE 8 HORAS APROXIMADAS
TURNOS AM
● Turno B1: Lunes a viernes 09:00 – 16:10, sábado 09:00 – 16:10. Domingo libre.
● Turno B2: Lunes a viernes 09:30 – 16:40, sábado 09:30 – 16:40. Domingo libre.
● Turno B3: Lunes a viernes 10:00 – 17:10, sábado 10:00 – 17:10. Domingo libre.
● Turno B4: Lunes a viernes 10:30 – 17:40, sábado 10:30 – 17:40. Domingo libre.
● Turno B5: Lunes a viernes 11:00 – 18:10, sábado 11:00 – 18:10. Domingo libre.
● Turno B6: Lunes a viernes 11:30 – 18:40, sábado 11:30 – 18:40. Domingo libre.
● Turno B7: Lunes a viernes 12:00 – 19:10, sábado 12:00 – 19:10. Domingo libre.
TURNOS PM
● Turno P1: Lunes a viernes 15:00 – 22:15, sábado 15:00 – 21:45. Domingo libre.
● Turno P2: Lunes a viernes 15:30 – 22:45, sábado 15:30 – 22:15. Domingo libre.
● Turno P3: Lunes a viernes 16:00 – 23:15, sábado 16:00 – 22:45. Domingo libre.
● Turno P4: Lunes a viernes 16:30 – 23:45, sábado 16:30 – 23:15. Domingo libre.
● Turno N1: Lunes a viernes 17:00 – 00:15, sábado 17:00 – 23:45. Domingo libre.
● Turno N2: Lunes a viernes 17:15 – 00:30, sábado 17:15 – 23:30. Domingo libre.
● Turno N3: Lunes a viernes 17:00 – 00:30, sábado 18:00 – 23:30. Domingo libre.
JORNADA DE 8,5 HORAS APROXIMADAS
TURNOS AM
● Turno C1: Lunes a jueves 09:00 – 17:30, viernes 09:00 – 17:30. Sábado y domingo libres.
● Turno C2: Lunes a jueves 09:30 – 18:00, viernes 09:30 – 18:00. Sábado y domingo libres.
● Turno C3: Lunes a jueves 10:00 – 18:30, viernes 10:00 – 18:30. Sábado y domingo libres.
● Turno C4: Lunes a jueves 10:30 – 19:00, viernes 10:30 – 19:00. Sábado y domingo libres.
● Turno C5: Lunes a jueves 11:00 – 19:30, viernes 11:00 – 19:30. Sábado y domingo libres.
● Turno C6: Lunes a jueves 11:30 – 20:00, viernes 11:30 – 20:00. Sábado y domingo libres.
TURNOS PM
● Turno P1: Lunes a jueves 15:00 – 22:15, viernes 15:00 – 22:15. Sábado y domingo libres.
● Turno P2: Lunes a jueves 15:30 – 22:45, viernes 15:30 – 22:45. Sábado y domingo libres.
● Turno P3: Lunes a jueves 16:00 – 23:15, viernes 16:00 – 23:15. Sábado y domingo libres.
● Turno P4: Lunes a jueves 16:30 – 23:45, viernes 16:30 – 23:45. Sábado y domingo libres.
● Turno N1: Lunes a jueves 17:00 – 00:15, viernes 17:00 – 00:15. Sábado y domingo libres.
● Turno N2: Lunes a jueves 17:30 – 00:30, viernes 17:30 – 00:30. Sábado y domingo libres.
● Turno N3: Lunes a jueves 18:00 – 00:30, viernes 18:00 – 00:30. Sábado y domingo libres.
JORNADA DE 9 HORAS APROXIMADAS
TURNOS AM
● Turno D1: Lunes a miércoles 09:00 – 19:00, jueves 09:00 – 18:00, viernes 09:00 – 13:00. Sábado y domingo libres.
● Turno D2: Lunes a miércoles 09:30 – 19:30, jueves 09:30 – 18:30, viernes 09:30 – 13:30. Sábado y domingo libres.
● Turno D3: Lunes a miércoles 10:00 – 20:00, jueves 10:00 – 19:00, viernes 10:00 – 14:00. Sábado y domingo libres.
● Turno D4: Lunes a miércoles 10:30 – 20:30, jueves 10:30 – 19:30, viernes 10:30 – 14:30. Sábado y domingo libres.
● Turno D5: Lunes a miércoles 11:00 – 21:00, jueves 11:00 – 20:00, viernes 11:00 – 15:00. Sábado y domingo libres.
● Turno D6: Lunes a miércoles 11:30 – 20:30, jueves 11:30 – 19:30, viernes 11:30 – 15:30. Sábado y domingo libres.
● Turno D7: Lunes a miércoles 12:00 – 21:00, jueves 12:00 – 20:00, viernes 12:00 – 16:00. Sábado y domingo libres.
TURNOS PM
● Turno P1: Lunes a miércoles 15:00 – 00:00, jueves 15:00 – 22:00, viernes 15:00 – 19:00. Sábado y domingo libres.
● Turno P2: Lunes a miércoles 15:30 – 00:30, jueves 15:30 – 22:30, viernes 15:30 – 19:30. Sábado y domingo libres.
● Turno P3: Lunes a miércoles 16:00 – 00:30, jueves 16:00 – 23:00, viernes 16:00 – 20:00. Sábado y domingo libres.
● Turno P4: Lunes a miércoles 16:30 – 00:30, jueves 16:30 – 23:30, viernes 16:30 – 20:30. Sábado y domingo libres.
JORNADA DE 9,5 HORAS APROXIMADAS
TURNOS AM
● Turno E1: Lunes a miércoles 09:00 – 19:00, jueves 09:00 – 18:00, viernes 09:00 – 13:00. Sábado y domingo libres.
● Turno E2: Lunes a miércoles 09:30 – 19:30, jueves 09:30 – 18:30, viernes 09:30 – 13:30. Sábado y domingo libres.
● Turno E3: Lunes a miércoles 10:00 – 20:00, jueves 10:00 – 19:00, viernes 10:00 – 14:00. Sábado y domingo libres.
● Turno E4: Lunes a miércoles 10:30 – 20:30, jueves 10:30 – 19:30, viernes 10:30 – 14:30. Sábado y domingo libres.
● Turno E5: Lunes a miércoles 11:00 – 21:00, jueves 11:00 – 20:00, viernes 11:00 – 15:00. Sábado y domingo libres.
● Turno E6: Lunes a miércoles 12:00 – 22:00, jueves 12:00 – 20:30, viernes 12:00 – 14:00. Sábado y domingo libres.
TURNOS PM
● Turno P1: Lunes a miércoles 13:00 – 23:00, jueves 13:00 – 22:00, viernes 13:00 – 17:00. Sábado y domingo libres.
● Turno P2: Lunes a miércoles 13:30 – 23:30, jueves 13:30 – 22:30, viernes 13:30 – 17:30. Sábado y domingo libres.
● Turno P3: Lunes a miércoles 14:00 – 00:00, jueves 14:00 – 23:00, viernes 14:00 – 18:00. Sábado y domingo libres.
● Turno P4: Lunes a miércoles 14:30 – 00:30, jueves 14:30 – 23:30, viernes 14:30 – 18:30. Sábado y domingo libres.
JORNADA DE 10 HORAS APROXIMADAS
TURNOS AM
● Turno F1: Lunes a miércoles 09:00 – 19:00, jueves 09:00 – 18:00, viernes 09:00 – 13:00. Sábado y domingo libres.
● Turno F2: Lunes a miércoles 09:30 – 19:30, jueves 09:30 – 18:30, viernes 09:30 – 13:30. Sábado y domingo libres.
● Turno F3: Lunes a miércoles 10:00 – 20:00, jueves 10:00 – 19:00, viernes 10:00 – 14:00. Sábado y domingo libres.
● Turno F4: Lunes a miércoles 10:30 – 20:30, jueves 10:30 – 19:30, viernes 10:30 – 14:30. Sábado y domingo libres.
● Turno F5: Lunes a miércoles 11:00 – 21:00, jueves 11:00 – 20:00, viernes 11:00 – 15:00. Sábado y domingo libres.
TURNOS PM
● Turno P1: Lunes a miércoles 13:00 – 23:00, jueves 13:00 – 22:00, viernes 13:00 – 17:00. Sábado y domingo libres.
● Turno P2: Lunes a miércoles 13:30 – 23:30, jueves 13:30 – 22:30, viernes 13:30 – 17:30. Sábado y domingo libres.
● Turno P3: Lunes a miércoles 14:00 – 00:00, jueves 14:00 – 23:00, viernes 14:00 – 18:00. Sábado y domingo libres.
● Turno P4: Lunes a miércoles 14:30 – 00:30, jueves 14:30 – 23:30, viernes 14:30 – 18:30. Sábado y domingo libres.
● Turno P5: Lunes a miércoles 13:00 – 21:00, jueves 13:00 – 20:30, viernes 13:00 – 17:00. Sábado y domingo libres.
● Turno P6: Lunes a miércoles 13:30 – 22:00, jueves 13:30 – 21:30, viernes 13:30 – 17:30. Sábado y domingo libres.
● Turno P7: Lunes a miércoles 14:00 – 22:30, jueves 14:00 – 21:30, viernes 14:00 – 18:00. Sábado y domingo libres.
● Turno N1: Lunes a miércoles 16:00 – 00:30, jueves 16:00 – 23:30, viernes 16:00 – 20:00. Sábado y domingo libres.
● Turno N2: Lunes a miércoles 16:30 – 00:30, jueves 16:30 – 23:30, viernes 16:30 – 20:30. Sábado y domingo libres.
● Turno N3: Lunes a miércoles 17:00 – 00:30, jueves 17:00 – 23:30, viernes 17:00 – 21:00. Sábado y domingo libres.
"""

CIERRE_TURNOS = (
    "Las jornadas y turnos señalados en el presente contrato tendrán carácter referencial. "
    "Además, deben ser pactados expresamente y comunicados al trabajador con una anticipación mínima de una semana, "
    "y podrán ajustarse según las necesidades operativas de la empresa, respetando siempre los límites legales de "
    "jornada y descanso establecidos en la legislación vigente. En todo caso, dichas modificaciones no podrán implicar "
    "menoscabo para el trabajador ni exceder los límites legales aplicables. Se deja establecido que los horarios de "
    "los turnos se encuentran en el Reglamento Interno de la empresa."
)

CLAUSULA_EXTRANJERO = [
    ("DÉCIMO PRIMERO: CLÁUSULAS ESPECIALES (TRABAJADOR EXTRANJERO). ", "Se establecen las siguientes:"),
    ("li", "- Vigencia: La obligación de prestar servicios emanada del presente contrato sólo podrá cumplirse una vez que el trabajador haya obtenido la visación de residencia correspondiente en Chile o el permiso especial de trabajo para extranjeros con visa en trámite."),
    ("li", "- Régimen previsional: El empleador se compromete a efectuar las retenciones correspondientes y entregarlas a las instituciones de seguridad social, salvo que las partes se acojan a la Ley 18.156."),
    ("li", "- Impuesto a la renta: El empleador se obliga a responder por el pago del impuesto a la renta correspondiente a la remuneración del trabajador extranjero, para rentas superiores a 13,5 UTM."),
]


def build(extranjero: bool, out_path: str):
    d = Document()
    d.styles['Normal'].paragraph_format.space_after = Pt(0)

    def cab(t):
        p = d.add_paragraph(); p.alignment = AL.CENTER; r = p.add_run(t); r.bold = True

    def vacio():
        d.add_paragraph()

    def cl(lead, resto=""):
        p = d.add_paragraph(); p.alignment = AL.JUSTIFY
        if lead:
            r = p.add_run(lead); r.bold = True
        if resto:
            p.add_run(resto)

    def li(t, bold=False):
        p = d.add_paragraph(); p.alignment = AL.JUSTIFY
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(t); r.bold = bold

    cab("CONTRATO INDIVIDUAL DE TRABAJO")
    cab("Ayudante de Cocina (Part Time)")
    vacio()
    cl("", 'En {CIUDAD_FIRMA}, a {FECHA_INICIO_CONTRATO}, entre {RAZON_SOCIAL}, RUT {RUT_EMPRESA}, representada por don {NOMBRE_REP_LEGAL}, Rut: {RUT_REP_LEGAL}, ambos domiciliados en Freire N° 1551, local 5, Quillota, Región de Valparaíso, correo electrónico {EMAIL_EMPRESA} que en adelante se denominará "empleador", y don {NOMBRE_EMPLEADO}, de nacionalidad {NACIONALIDAD_EMPLEADO}, Rut {RUT_EMPLEADO}, fecha de nacimiento {FECHA_NACIMIENTO}, de estado civil {ESTADO_CIVIL}, con domicilio {DIRECCION_EMPLEADO}, comuna de {COMUNA_EMPLEADO}, correo electrónico {EMAILPERSONAL_EMPLEADO} que en adelante se denominará "el trabajador", se ha convenido en el siguiente contrato de trabajo.')
    vacio()
    cl("PRIMERO: NATURALEZA DE LOS SERVICIOS Y FUNCIONES. ", "El Trabajador se obliga a prestar servicios en el cargo de AYUDANTE DE COCINA, desempeñando labores de producción, higiene y apoyo operativo en la elaboración de alimentos.")
    cl("1. Funciones Principales y Responsabilidades: ", "El Trabajador deberá apoyar la operación integral de la cocina bajo la supervisión del Jefe de Cocina. Sus labores específicas incluyen:")
    li('- Producción y Mise en Place: realizar la preparación previa de ingredientes (cortar, picar, marinar) y producción según recetario oficial.')
    li('- Operación de Cuartos: ejercer labores de cocina tanto en cuarto frío como en cuarto caliente, asegurando la correcta cocción y montaje.')
    li('- Logística Operativa: ejercer labores de "tercero" o runner, recibiendo comandas y entregando el producto terminado al área de caja/reparto.')
    li('- Orden e Higiene: mantener la limpieza profunda de la cocina, utensilios y equipos. Almacenar alimentos en bodega siguiendo estrictamente las prácticas sanitarias (FIFO) y de seguridad.')
    li('- Apoyo al Equipo: colaborar en la inducción de nuevos ingresos y asumir la encargatura de turno si es solicitado por la jefatura.')
    cl("2. Otras funciones relevantes: ", "El Trabajador deberá estar capacitado para cubrir cualquier estación de trabajo en caso de emergencia y realizar labores administrativas conexas al cargo.")
    vacio()
    cl("SEGUNDO: JORNADA DE TRABAJO. ", "La jornada de trabajo será de 30 horas semanales, distribuida de lunes a domingo conforme al sistema de turnos rotativos y variables que el Empleador publicará semanalmente, con un día de descanso a la semana (y al menos dos domingos libres al mes). El Trabajador contará con una colación de 30 minutos diarios (no imputables a la jornada). De lo anterior, se seguirá lo informado en el respectivo Reglamento Interno de Orden, Higiene y Seguridad. Los turnos de trabajo, de carácter referencial, se distribuirán conforme a las siguientes tablas:")
    vacio()
    add_turnos(d)
    cl("", CIERRE_TURNOS)
    vacio()
    cl("TERCERO: LUGAR DE TRABAJO. ", "Preferentemente en Calle Ramón Freire 1551, Local 5, Quillota. El Empleador podrá trasladar al Trabajador a otras sucursales dentro de la misma ciudad, en ejercicio de la facultad de Ius Variandi, sin que ello signifique menoscabo.")
    vacio()
    cl("CUARTO: REMUNERACIÓN. ", "El Empleador pagará al Trabajador las siguientes prestaciones mensuales:")
    li("1) Sueldo Base: $ {SUELDO_BASE} ({SUELDO_BASE_PALABRA})")
    li("2) Asignación de movilización: $ {ASIGNACION_MOVILIZACION} ({ASIGNACION_MOVILIZACION_PALABRA})")
    li("3) Asignación de colación: $ {ASIGNACION_COLACION} ({ASIGNACION_COLACION_PALABRA})")
    cl("", "{GRATIFICACION_TEXTO}")
    cl("", "El pago de las remuneraciones se realizará preferentemente por medio de transferencia electrónica a la cuenta informada por el Trabajador.")
    vacio()
    cl("QUINTO: ", "Las remuneraciones antes referidas quedarán afectas a los descuentos previsionales y tributarios que correspondan y a los descuentos legales pertinentes. El Trabajador acepta y autoriza, desde ya, que pueda descontarse de su remuneración el tiempo no trabajado, sea por permisos, atrasos, faltas, rotura de implementos u otros motivos, como también por anticipos de remuneración solicitados por él.")
    vacio()
    cl("SEXTO: OBLIGACIONES Y PROHIBICIONES ESPECIALES (COCINA). ", "Se establecen como obligaciones esenciales, cuyo incumplimiento se considerará grave:")
    li("1. Normas Sanitarias: es causal de término inmediato el incumplimiento de normas de higiene que ponga en riesgo la salud de los clientes (ej.: contaminación cruzada, falta de aseo personal).")
    li("2. Mermas y Consumo: queda prohibido el consumo de insumos o productos terminados sin autorización y el desperdicio injustificado de materias primas.")
    li("3. Seguridad: es obligatorio el uso de EPP (cofia, zapatos antideslizantes, guantes) en todo momento dentro del área de producción.")
    vacio()
    cl("SÉPTIMO: CONFIDENCIALIDAD. ", "El Trabajador guardará estricta reserva sobre las recetas, fichas técnicas, proveedores y procedimientos exclusivos de {RAZON_SOCIAL} (secreto industrial).")
    vacio()
    cl("OCTAVO: PREVISIÓN Y SALUD. ", "Se deja constancia que el Trabajador cotizará en AFP {PREVISION} y en el sistema de salud {INSTITUCION_SALUD}.")
    vacio()
    cl("NOVENO: COMUNICACIONES ELECTRÓNICAS. ", "El Trabajador autoriza expresamente recibir su documentación laboral (liquidaciones, comprobantes de feriado, anexos) y notificaciones al correo electrónico personal: {EMAILPERSONAL_EMPLEADO}. El Empleador fija como correo oficial: {EMAIL_EMPRESA}.")
    vacio()
    cl("DÉCIMO: VIGENCIA. ", "El presente contrato tiene una duración fija, reconociendo la fecha de ingreso actual a esta empresa el día {FECHA_INICIO_CONTRATO} y concluyendo la relación laboral el día {FECHA_TERMINO_CONTRATO}, sin perjuicio de las prórrogas que de común acuerdo puedan alcanzar las partes si así lo determinaran.")
    vacio()
    if extranjero:
        for item in CLAUSULA_EXTRANJERO:
            if item[0] == "li":
                li(item[1])
            else:
                cl(item[0], item[1])
        vacio()
        n_jur, n_ej = "DÉCIMO SEGUNDO", "DÉCIMO TERCERO"
    else:
        n_jur, n_ej = "DÉCIMO PRIMERO", "DÉCIMO SEGUNDO"
    cl(f"{n_jur}: JURISDICCIÓN. ", "Para todos los efectos legales derivados de este contrato, las partes fijan su domicilio en la ciudad de Quillota, sometiéndose a la jurisdicción y competencia de sus Tribunales de Justicia.")
    vacio()
    cl(f"{n_ej}: EJEMPLARES. ", "Se suscribe este instrumento en dos ejemplares de igual tenor, quedando uno de ellos en poder del empleador y el restante en poder del trabajador, quien declara recibirlo en este acto.")
    vacio(); vacio()
    tb = d.add_table(rows=3, cols=2); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (a, b) in enumerate([("{RAZON_SOCIAL}", "{NOMBRE_EMPLEADO}"), ("RUT {RUT_EMPRESA}", "RUT {RUT_EMPLEADO}"), ("EMPLEADOR", "TRABAJADOR")]):
        c = tb.rows[i].cells
        for cell, txt in ((c[0], a), (c[1], b)):
            cell.paragraphs[0].alignment = AL.CENTER; cell.paragraphs[0].add_run(txt)
    d.save(out_path)
    print("Generado:", out_path, "| párrafos:", len(d.paragraphs))


def build_anexo(out_path):
    """Anexo de jornada laboral: empresa pre-llenada, datos del trabajador y
    turno asignado en blanco para completar a mano. Form impreso (sin placeholders)."""
    d = Document(); d.styles['Normal'].paragraph_format.space_after = Pt(0)
    L = "______________________"

    def cab(t):
        p = d.add_paragraph(); p.alignment = AL.CENTER; p.add_run(t).bold = True

    def vacio():
        d.add_paragraph()

    def cl(lead, resto=""):
        p = d.add_paragraph(); p.alignment = AL.JUSTIFY
        if lead:
            p.add_run(lead).bold = True
        if resto:
            p.add_run(resto)

    cab("ANEXO DE CONTRATO INDIVIDUAL DE TRABAJO")
    cab("(Establecimiento de Jornada Laboral)")
    vacio()
    cl("", 'En Quillota, a ____ de ' + L + ' de 2026, entre JEREZ DE LA FRONTERA SPA, RUT 78.269.062-0, representada legalmente por don Daniel Elías Améstica Hernández, RUT 13.635.853-7, ambos domiciliados en Avenida Freire 1551, Lote A, Depto. #L5, Quillota, Región de Valparaíso, correo electrónico danielamesticah@gmail.com, en adelante "el empleador", y don(ña) _________________________________________________, cédula de identidad N° ____________________, en adelante "el trabajador", se ha convenido el siguiente anexo al contrato individual de trabajo suscrito con fecha ____ de ' + L + ' de __________:')
    vacio()
    cl("PRIMERO. ", "Las partes dejan constancia de que el contrato individual de trabajo vigente no especificaba la distribución de la jornada de trabajo. Por el presente acto, de común acuerdo, regularizan y establecen expresamente la jornada laboral del trabajador conforme a las cláusulas siguientes.")
    vacio()
    cl("SEGUNDO. ", "La jornada ordinaria de trabajo será de ________ horas semanales, distribuida según el sistema de turnos rotativos y variables que el empleador publica semanalmente, con los descansos legales correspondientes.")
    vacio()
    cl("", "Para todos los efectos, los turnos referenciales de la empresa son los siguientes:")
    vacio()
    add_turnos(d)
    cl("", CIERRE_TURNOS)
    vacio()
    cl("TERCERO. ", "El presente anexo rige a partir del día ____ de " + L + " de __________. En todo lo no modificado por este instrumento, se mantienen vigentes e inalteradas todas las demás estipulaciones del contrato individual de trabajo.")
    vacio()
    cl("CUARTO. ", "Se firma en dos ejemplares de igual tenor y fecha, quedando uno en poder de cada parte.")
    vacio(); vacio()
    tb = d.add_table(rows=5, cols=2); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    filas = [
        ("_______________________________", "_______________________________"),
        ("JEREZ DE LA FRONTERA SPA", "Nombre: ______________________"),
        ("RUT 78.269.062-0", "RUT: ______________________"),
        ("p.p. Daniel Elías Améstica Hernández", ""),
        ("EMPLEADOR", "TRABAJADOR"),
    ]
    for i, (a, b) in enumerate(filas):
        c = tb.rows[i].cells
        for cell, txt in ((c[0], a), (c[1], b)):
            cell.paragraphs[0].alignment = AL.CENTER; cell.paragraphs[0].add_run(txt)
    d.save(out_path)
    print("Generado:", out_path, "| párrafos:", len(d.paragraphs))


BASE = "plantillas/JEREZ 78269062-0/"
build(False, BASE + "CONTRATO Ayudante Cocina PT - Plazo Fijo.docx")
build(True, BASE + "CONTRATO Ayudante Cocina PT Extranjero - Plazo Fijo.docx")
build_anexo(BASE + "ANEXO Jornada Laboral (form manual).docx")
