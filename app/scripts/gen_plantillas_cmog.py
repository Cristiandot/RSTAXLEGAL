# -*- coding: utf-8 -*-
"""
Generador de plantillas de contrato de GRUPO C MOG SPA (78.266.284-8) —
cliente C.21 Hermanos Ormeño, franquicia del restaurante de JEREZ.
Base: contrato tipo "Gerente de Local" entregado por Cristian (jul-2026),
adaptado desde el de JEREZ DE LA FRONTERA (se quitó el reconocimiento de
antigüedad LCA Gastronomía y los datos hardcodeados de Quillota).

Cargo excluido de la limitación de jornada (sin registro de asistencia ni
horas extra), por lo que la plantilla no usa {HORAS_SEMANALES}: una sola
plantilla sirve plazo fijo/indefinido vía {TIPO_CONTRATO} y
({FECHA_TERMINO_CONTRATO}).

Ejecutar:  py scripts/gen_plantillas_cmog.py   (desde app/)
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "plantillas/GRUPO C MOG 78266284-8/"
OUT_GERENTE = BASE + "CONTRATO Gerente de Local.docx"
OUT_GERENTE_EXT = BASE + "CONTRATO Gerente de Local Extranjero.docx"

# Cláusula obligatoria en toda plantilla de nacionalidad extranjero (regla RSTL).
CLAUSULA_EXTRANJERO = [
    ("head", "Se establecen las siguientes:"),
    ("li", "- Vigencia: La obligación de prestar servicios emanada del presente contrato sólo podrá cumplirse una vez que el trabajador haya obtenido la visación de residencia correspondiente en Chile o el permiso especial de trabajo para extranjeros con visa en trámite."),
    ("li", "- Régimen previsional: El empleador se compromete a efectuar las retenciones correspondientes y entregarlas a las instituciones de seguridad social, salvo que las partes se acojan a la Ley 18.156."),
    ("li", "- Impuesto a la renta: El empleador se obliga a responder por el pago del impuesto a la renta correspondiente a la remuneración del trabajador extranjero, para rentas superiores a 13,5 UTM."),
]

ORDINALES = [
    "PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO", "SEXTO",
    "SÉPTIMO", "OCTAVO", "NOVENO", "DÉCIMO", "DÉCIMO PRIMERO",
    "DÉCIMO SEGUNDO", "DÉCIMO TERCERO", "DÉCIMO CUARTO",
]


def build_contrato(out_path: str, extranjero: bool = False):
    d = Document()
    d.styles["Normal"].paragraph_format.space_after = Pt(0)

    def cab(t):
        p = d.add_paragraph(); p.alignment = AL.CENTER; r = p.add_run(t); r.bold = True

    def vacio():
        d.add_paragraph()

    def cl(lead, resto=""):
        p = d.add_paragraph(); p.alignment = AL.JUSTIFY
        if lead:
            r = p.add_run(lead); r.bold = True
        if resto:
            p.add_run(resto)

    def num(t, bold=False):
        # ítem numerado (el número va en el texto), con sangría colgante
        p = d.add_paragraph(); p.alignment = AL.JUSTIFY
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(t); r.bold = bold

    n = iter(ORDINALES)

    cab("CONTRATO INDIVIDUAL DE TRABAJO")
    cab("{CARGO}")
    vacio()
    cl("", 'En {CIUDAD_FIRMA}, a {FECHA_INICIO_CONTRATO}, entre {RAZON_SOCIAL}, RUT {RUT_EMPRESA}, representada por don {NOMBRE_REP_LEGAL}, Rut: {RUT_REP_LEGAL}, ambos domiciliados en {DIRECCION_EMPRESA}, comuna de {COMUNA_EMPRESA}, correo electrónico {EMAIL_EMPRESA}, que en adelante se denominará "el empleador", y don(ña) {NOMBRE_EMPLEADO}, de nacionalidad {NACIONALIDAD_EMPLEADO}, Rut {RUT_EMPLEADO}, fecha de nacimiento {FECHA_NACIMIENTO}, de estado civil {ESTADO_CIVIL}, con domicilio {DIRECCION_EMPLEADO}, comuna de {COMUNA_EMPLEADO}, correo electrónico {EMAILPERSONAL_EMPLEADO}, que en adelante se denominará "el trabajador", se ha convenido en el siguiente contrato de trabajo:')
    vacio()
    cl(f"{next(n)}: NATURALEZA DE LOS SERVICIOS Y FUNCIONES. ", "El Trabajador asume el cargo de confianza de {CARGO}, obligándose a la administración integral de la unidad de negocio, liderando los equipos de trabajo y velando por la rentabilidad y operatividad del restaurante.")
    cl("", "1. Funciones Principales y Responsabilidades: El Trabajador deberá coordinar las operaciones diarias, asegurando la interacción fluida entre cocina, caja y salón. Sus labores específicas incluyen:")
    num("- Liderazgo y Gestión de Personas: Dirigir al personal de caja y cocina, encargándose del reclutamiento, inducción, capacitación, creación de horarios/turnos y supervisión de la distribución de propinas según la política interna.")
    num("- Gestión Operativa y de Inventarios: Velar por el orden y limpieza de todas las áreas (salón, bodega, terraza, oficina). Recibir proveedores, ingresar facturas al sistema y mantener el stock crítico, reportando necesidades de abastecimiento.")
    num("- Control Financiero y KPIs: Asegurar el cumplimiento de las metas de venta, control de costos de nómina y compras. Responsable final de la apertura, cierre y cuadratura de cajas, así como de la gestión de depósitos bancarios si corresponde.")
    num("- Servicio al Cliente: Gestionar y resolver problemas o reclamos de clientes (escalamiento), asegurando la imagen de la marca y la experiencia de servicio.")
    num("- Otras funciones relevantes: Dada la naturaleza del rubro gastronómico, el Trabajador acepta ejecutar labores conexas o complementarias a las principales, tales como apoyo operativo en salón o empaque, cuando las necesidades del servicio lo requieran.")
    vacio()
    cl(f"{next(n)}: JORNADA DE TRABAJO. ", "Dada la naturaleza de su cargo, que implica facultades de administración y supervisión sin fiscalización superior inmediata, el Trabajador queda excluido de la limitación de jornada de trabajo, de conformidad con la legislación vigente. En consecuencia, no estará obligado a registrar asistencia y no tendrá derecho al pago de horas extraordinarias. Sin perjuicio de lo anterior, el Trabajador distribuirá su tiempo de manera eficiente para cumplir con los objetivos del cargo y la supervisión de los turnos de funcionamiento del local.")
    vacio()
    cl(f"{next(n)}: LUGAR DE TRABAJO. ", "El Trabajador prestará sus servicios preferentemente en {DIRECCION_EMPRESA}, comuna de {COMUNA_EMPRESA}. El Empleador podrá trasladar al Trabajador a otras sucursales dentro de la misma ciudad, sin que ello signifique menoscabo para el Trabajador.")
    vacio()
    cl(f"{next(n)}: REMUNERACIÓN. ", "El Empleador pagará al Trabajador las siguientes prestaciones mensuales:")
    num("1) Sueldo Base: $ {SUELDO_BASE} ({SUELDO_BASE_PALABRA}).")
    num("2) Asignación de movilización: $ {ASIGNACION_MOVILIZACION} ({ASIGNACION_MOVILIZACION_PALABRA}).")
    num("3) Asignación de Colación: $ {ASIGNACION_COLACION} ({ASIGNACION_COLACION_PALABRA}).")
    cl("", "{GRATIFICACION_TEXTO}")
    cl("", "El pago de las remuneraciones se realizará preferentemente por medio de transferencia electrónica a la cuenta informada por el Trabajador.")
    vacio()
    cl(f"{next(n)}: ", "Las remuneraciones antes referidas quedarán afectas a los descuentos previsionales y tributarios que correspondan y a los demás que autorice la ley. El trabajador acepta y autoriza, desde ya, que pueda descontarse de su remuneración el tiempo no trabajado, sea por permisos, atrasos, faltas, rotura de implementos, u otros motivos, como también por anticipos de remuneración solicitados por él.")
    vacio()
    cl(f"{next(n)}: OBLIGACIONES Y PROHIBICIONES ESPECIALES (CARGO DE CONFIANZA). ", "Considerando el cargo de {CARGO}, se establecen como obligaciones esenciales, cuyo incumplimiento se considerará incumplimiento grave de las obligaciones que impone el contrato:")
    num("1. Cumplimiento de KPIs: Es obligación esencial velar por las metas de venta y control de costos (mermas/nómina) definidas por la Gerencia General.")
    num("2. Gestión de Valores: Supervisar que los cierres de caja de su equipo sean fidedignos. Cualquier inconsistencia financiera en el local bajo su cargo será objeto de investigación y responsabilidad administrativa.")
    num("3. Clima Laboral y Trato: Mantener un trato respetuoso y profesional. Queda prohibido el abuso de autoridad o la falta de equidad en la asignación de turnos o distribución de propinas.")
    num("4. Imagen y Mantenimiento: Es su responsabilidad directa reportar y gestionar la solución de fallas de infraestructura o equipos que afecten la operación.")
    vacio()
    cl(f"{next(n)}: CONFIDENCIALIDAD. ", "El Trabajador guardará absoluta reserva sobre los procedimientos de caja, claves de acceso al sistema, datos de proveedores, recetas y montos de recaudación de la empresa.")
    vacio()
    cl(f"{next(n)}: PREVISIÓN Y SALUD. ", "Se deja constancia que el Trabajador cotizará en {PREVISION} y en el sistema de salud {INSTITUCION_SALUD}.")
    vacio()
    cl(f"{next(n)}: COMUNICACIONES ELECTRÓNICAS. ", "El Trabajador autoriza expresamente recibir su documentación laboral (liquidaciones, comprobantes de feriado, anexos) y notificaciones al correo electrónico personal: {EMAILPERSONAL_EMPLEADO}. El Empleador fija como correo oficial: {EMAIL_EMPRESA}.")
    vacio()
    cl(f"{next(n)}: VIGENCIA. ", "La duración de este contrato será {TIPO_CONTRATO} ({FECHA_TERMINO_CONTRATO}), rigiendo para su término las disposiciones establecidas en la legislación vigente. Se deja constancia que el trabajador ingresó a prestar servicios para el empleador con fecha {FECHA_INICIO_CONTRATO}.")
    vacio()
    if extranjero:
        cl(f"{next(n)}: CLÁUSULAS ESPECIALES (TRABAJADOR EXTRANJERO). ", CLAUSULA_EXTRANJERO[0][1])
        for _, texto in CLAUSULA_EXTRANJERO[1:]:
            num(texto)
        vacio()
    cl("", "{CLAUSULAS_ADICIONALES}")
    vacio()
    cl(f"{next(n)}: ", "Para todos los efectos legales derivados de este contrato, las partes fijan su domicilio en la ciudad de {CIUDAD_FIRMA}, sometiéndose a la jurisdicción y competencia de sus Tribunales de Justicia.")
    vacio()
    cl(f"{next(n)}: ", "Se suscribe este instrumento en dos ejemplares de igual tenor, quedando uno de ellos en poder del empleador y el restante en poder del trabajador, quien declara recibirlo en este acto.")
    vacio(); vacio()
    tb = d.add_table(rows=3, cols=2); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    filas = [("{RAZON_SOCIAL}", "{NOMBRE_EMPLEADO}"),
             ("RUT {RUT_EMPRESA}", "RUT {RUT_EMPLEADO}"),
             ("EMPLEADOR", "TRABAJADOR")]
    for i, (a, b) in enumerate(filas):
        c = tb.rows[i].cells
        for cell, txt in ((c[0], a), (c[1], b)):
            cell.paragraphs[0].alignment = AL.CENTER
            cell.paragraphs[0].add_run(txt)
    d.save(out_path)
    print("Generado:", out_path, "| párrafos:", len(d.paragraphs))


if __name__ == "__main__":
    os.makedirs(BASE, exist_ok=True)
    build_contrato(OUT_GERENTE)
    build_contrato(OUT_GERENTE_EXT, extranjero=True)
